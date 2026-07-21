from io import BytesIO

from docx import Document
from flask import Blueprint, jsonify, request
from PyPDF2 import PdfReader

from ai.cv_parser import parse_cv_text
from ai.recommender import recommend_jobs_for_skills
from database.models import get_jobs

upload_bp = Blueprint("upload", __name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
        if sum(len(p) for p in parts) > 200_000:
            break
    return "\n".join(parts).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Also pull simple table cell text when present.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts).strip()


@upload_bp.post("/cv/upload")
def upload_cv():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Missing file."}), 400

    filename = (file.filename or "").lower()
    if not any(filename.endswith(ext) for ext in ALLOWED_EXTENSIONS):
        return jsonify({"error": "Only PDF and DOCX files are supported."}), 400

    file_bytes = file.read()
    if not file_bytes:
        return jsonify({"error": "Empty file."}), 400

    try:
        if filename.endswith(".pdf"):
            cv_text = _extract_pdf_text(file_bytes)
        else:
            cv_text = _extract_docx_text(file_bytes)
    except Exception as e:
        return jsonify({"error": f"Failed to read file: {e}"}), 400

    if not cv_text:
        return jsonify({"error": "Could not extract text from the CV."}), 400

    cv_profile = parse_cv_text(cv_text)
    skills = cv_profile["skills"]
    cv_field = cv_profile["field"]

    jobs_rows = get_jobs(
        keyword=None,
        location=None,
        limit=5000,
        exclude_sources=["manual"],
    )

    try:
        recommendations = recommend_jobs_for_skills(
            skills=skills,
            jobs=jobs_rows,
            cv_text=cv_text,
            cv_field=cv_field or None,
            cv_experience=(cv_profile.get("experience") or [None])[0],
            top_k=20,
        )
    except RuntimeError as error:
        return jsonify({"error": str(error)}), 503

    return jsonify(
        {
            "skills": skills,
            "cv_field": cv_field,
            "cv_position": cv_profile["position"],
            "cv_profile": cv_profile,
            "recommendations": recommendations,
            "jobs_considered": len(jobs_rows),
        }
    )
